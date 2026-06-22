#!/usr/bin/env python3
"""
Generate reviewer-guide.docx — a Word document for academic advisors reviewing
the scraped ISU Plans of Study program data before publication.

Usage:
    python3 scripts/generate-reviewer-guide.py
    python3 scripts/generate-reviewer-guide.py --out path/to/guide.docx
"""

import argparse
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── ISU brand colours ─────────────────────────────────────────────────────────
ISU_RED    = RGBColor(0xC0, 0x00, 0x00)
ISU_GREY   = RGBColor(0x59, 0x59, 0x59)
LIGHT_RED  = RGBColor(0xF9, 0xEB, 0xEB)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)


# ── helpers ───────────────────────────────────────────────────────────────────

def set_table_full_width(table):
    tbl   = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW  = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")   # 5000 = 100 % in Word's 1/50-percent units
    tblPr.append(tblW)


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(rgb))
    tcPr.append(shd)


def heading(doc, text, level=1, colour=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if colour:
        run.font.color.rgb = colour
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def callout(doc, title, text, bg=LIGHT_RED, title_colour=ISU_RED):
    """A single-cell bordered table used as a callout / note box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Inches(6)

    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title + "  ")
    r1.bold = True
    r1.font.color.rgb = title_colour
    r1.font.size = Pt(10)
    r2 = p1.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = ISU_GREY

    # border the table
    tbl_pr = tbl._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), str(title_colour))
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)
    set_table_full_width(tbl)
    doc.add_paragraph()


def two_col_table(doc, rows, headers=None, col_widths=None):
    """Generic bordered table."""
    n_cols = len(rows[0])
    tbl = doc.add_table(rows=0, cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    if headers:
        hdr_row = tbl.add_row()
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            set_cell_bg(cell, ISU_RED)
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9)
            if col_widths:
                cell.width = Inches(col_widths[i])

    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        bg = LIGHT_GREY if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if col_widths:
                cell.width = Inches(col_widths[ci])
    set_table_full_width(tbl)
    doc.add_paragraph()


# ── document sections ─────────────────────────────────────────────────────────

def section_introduction(doc):
    heading(doc, "Introduction", level=1, colour=ISU_RED)
    body(doc,
        "Illinois State University is replacing hand-produced paper advising worksheets with "
        "generated, interactive documents for transfer students. This project produces two "
        "formats — a fillable PDF and an interactive web worksheet — for every undergraduate "
        "degree program at ISU, covering approximately 300 degree sequences across eight colleges."
    )
    body(doc,
        "Both formats are generated automatically from a shared data file for each program. "
        "Those data files were built by a software tool that reads the ISU undergraduate catalog, "
        "but some details require human judgment to verify or complete. This guide explains what "
        "that tool produced, what it cannot determine on its own, and exactly what we are asking "
        "reviewers to check."
    )
    callout(doc,
        "Your role:",
        "You are not being asked to rewrite the curriculum or the catalog. You are being asked "
        "to confirm that what the tool captured matches what the catalog actually says, and to "
        "flag anything that looks wrong or incomplete.",
        bg=LIGHT_RED, title_colour=ISU_RED,
    )


def section_data_collection(doc):
    heading(doc, "How the Data Were Collected", level=1, colour=ISU_RED)

    heading(doc, "The Authoritative Source", level=2)
    body(doc,
        "All program data come from the ISU Undergraduate Catalog "
        "(catalog.illinoisstate.edu). The catalog is the definitive record of degree "
        "requirements. Every program JSON file includes a direct link to its catalog page, "
        "which reviewers should consult when verifying any item."
    )

    heading(doc, "The CourseDog API", level=2)
    body(doc,
        "ISU manages its curriculum in a system called CourseDog. An internal API provides "
        "structured data about programs, courses, and general education attributes. The scraping "
        "tool queries this API to retrieve:"
    )
    for item in [
        "The list of required and elective courses for each program sequence.",
        "Credit hours for each course.",
        "General education (gen-ed) attribute codes that indicate which ISU or IAI requirement "
        "a course can satisfy (for example, GE14-MAT indicates Mathematics, GE14-SS indicates "
        "Social Sciences).",
        "Program metadata: degree type, college, department, minimum credit hours.",
    ]:
        bullet(doc, item)
    body(doc,
        "The tool then transforms this raw API data into a structured JSON file — one per "
        "program sequence — that serves as the single source of truth for both the HTML and "
        "PDF worksheets."
    )

    heading(doc, "What the Tool Can and Cannot Determine", level=2)
    body(doc,
        "The CourseDog API provides reliable data for straightforward requirements. "
        "However, several structural features of degree programs cannot be reliably inferred "
        "from the API alone and require human review:"
    )
    two_col_table(doc,
        rows=[
            ("Fixed required courses",
             "Captured automatically from the program's required course list."),
            ("Choose-one options (e.g., MAT 120 or MAT 121)",
             "Partially captured; inline choices embedded in fixed lists may be missed."),
            ("Elective pools (pick N from a list)",
             "Captured when the catalog groups them explicitly; free-form lists require review."),
            ("Gen-ed attribute mappings",
             "Captured via CourseDog attribute codes (GE14-*). A small number of courses "
             "carry unexpected codes that need catalog verification."),
            ("Phased program structure (e.g., pre-nursing then nursing core)",
             "Not captured — requires human annotation to split into phases."),
            ("Repeat/accumulate requirements (e.g., applied music, ensembles)",
             "Not captured — credit ranges and level-progression rules need manual entry."),
            ("Whether a course satisfies vs. presupposes a gen-ed requirement",
             "Not determinable from the API — this is the most common review task."),
        ],
        headers=["What the tool captures", "Reliability"],
        col_widths=[2.8, 3.4],
    )

    heading(doc, "The Seven Reference Programs", level=2)
    body(doc,
        "Before the full scrape, seven programs were hand-crafted by the development team as "
        "reference implementations. These cover the full range of structural complexity — "
        "phased programs, repeatable courses, escrow credit, compliance requirements, and more. "
        "They are retained in the system as the baseline for testing and comparison."
    )
    for item in [
        "Accountancy — Financial Accounting Sequence (B.S.)",
        "Art — Art History Sequence (B.A.)",
        "Early Childhood Education — Pedagogy Sequence (B.S.)",
        "Music — Composition/Theory Emphasis (B.M.)",
        "Nursing — Traditional Prelicensure Sequence (B.S.N.)",
        "Nursing — R.N. to B.S.N. Sequence (B.S.N.)",
        "Physics — Physics Teacher Education Sequence (B.S.)",
    ]:
        bullet(doc, item)


def section_worksheets(doc):
    heading(doc, "How the Data Feed into the Worksheets", level=1, colour=ISU_RED)

    heading(doc, "The JSON File as the Single Source of Truth", level=2)
    body(doc,
        "Each program has exactly one JSON data file. Both the HTML worksheet and the PDF "
        "worksheet are generated directly from that file — there is no separate authoring "
        "step. A change to the JSON is automatically reflected in both outputs the next time "
        "the build tools are run. This means that if a reviewer identifies an error, it is "
        "fixed once in the JSON and both documents are corrected."
    )

    heading(doc, "The HTML Worksheet", level=2)
    body(doc,
        "The HTML worksheet is an interactive web page that students and advisors use during "
        "advising sessions. Its key features are:"
    )
    for item in [
        ("Two-column layout: ", "General Education requirements appear on the left; major requirements appear on the right."),
        ("Track toggle: ", "Students select which general education track applies to them — ISU Gen Ed, IAI Transferable Core, or Completed Associate's Degree — and the left column updates to show only the relevant requirements."),
        ("Checkboxes and transfer fields: ", "Each requirement row has a checkbox and a text field for entering a transfer course equivalent."),
        ("Live progress totals: ", "Credit hour counts update as rows are checked, showing progress toward minimums."),
        ("Cross-reference propagation: ", "If a course satisfies both a major requirement and a gen-ed requirement, checking one automatically checks the other."),
    ]:
        bullet(doc, item[1], bold_prefix=item[0])

    heading(doc, "The PDF Worksheet", level=2)
    body(doc,
        "The PDF worksheet is a fillable AcroForm document intended for printing or "
        "electronic completion. It mirrors the HTML layout in portrait letter format and "
        "is produced in three versions per program — one for each gen-ed track. "
        "It uses standard PDF form fields (checkboxes and text boxes) and contains no "
        "scripting, making it compatible with any PDF viewer."
    )

    heading(doc, "The Three General Education Tracks", level=2)
    body(doc,
        "ISU accepts transfer students under three different general education frameworks. "
        "Every worksheet accommodates all three:"
    )
    two_col_table(doc,
        rows=[
            ("ISU Gen Ed",
             "Students who have not completed a transferable gen-ed core. They must satisfy "
             "ISU's own general education program (11 categories, ~39 credit hours)."),
            ("IAI Transferable Core",
             "Students who completed the Illinois Articulation Initiative (IAI) General "
             "Education Core Curriculum at an Illinois community college. The IAI core is "
             "accepted in place of ISU Gen Ed, with some additional requirements."),
            ("Completed Associate's Degree",
             "Students who hold a baccalaureate-oriented A.A. or A.S. degree from an "
             "Illinois institution. The associate's degree fulfills ISU Gen Ed in full."),
        ],
        headers=["Track", "Who it applies to"],
        col_widths=[1.8, 4.4],
    )


def section_workbook(doc):
    heading(doc, "How to Use the Review Workbook", level=1, colour=ISU_RED)
    body(doc,
        "Along with this guide you have received an Excel file: advisor-review.xlsx. "
        "It contains two sheets."
    )

    heading(doc, "Sheet 1 — Program Review", level=2)
    body(doc,
        "This sheet has one row per program sequence (300 rows total). Use it to track "
        "whether each program has been reviewed and approved."
    )
    two_col_table(doc,
        rows=[
            ("Program ID", "Unique identifier for the program sequence. Click to open the draft HTML worksheet on the development site."),
            ("Title / Sequence / Degree", "The program name and degree type from the catalog."),
            ("College", "The college responsible for the program. The sheet is pre-sorted by college so you can focus on your area."),
            ("Catalog", "Click to open the program's page in the ISU Undergraduate Catalog."),
            ("Reviewer", "Enter your name here when you take ownership of a program."),
            ("Status", "Use the dropdown: Pending (not yet reviewed), Approved (correct as-is), Changes Needed (issues found — document in Notes or Issue Log)."),
            ("Notes", "Free text for brief comments. For detailed issues, add a row to Sheet 2."),
        ],
        headers=["Column", "Purpose"],
        col_widths=[1.5, 4.7],
    )
    callout(doc,
        "Tip:",
        "Use the College column filter to show only your programs. "
        "The Program ID hyperlinks will work once the worksheets are deployed to the development site.",
        bg=LIGHT_GREY, title_colour=ISU_GREY,
    )

    heading(doc, "Sheet 2 — Issue Log", level=2)
    body(doc,
        "This sheet is pre-populated with items that need advisor judgment. There are two "
        "categories of pre-populated items:"
    )
    bullet(doc,
        "Structural gaps (8 items): specific programs where the tool could not fully encode "
        "the program structure from the catalog.",
        bold_prefix="Structural gaps — ",
    )
    bullet(doc,
        "Auto-fulfilled items (1,118 items): cases where one or more major courses appear to "
        "satisfy a general education requirement. Each one needs a human to confirm whether "
        "the course truly satisfies the requirement or merely presupposes it.",
        bold_prefix="Auto-fulfilled items — ",
    )
    body(doc,
        "You can also add new rows to this sheet when you find issues during your review."
    )
    two_col_table(doc,
        rows=[
            ("Program ID", "The program the issue belongs to."),
            ("College", "Use to filter to your college's items."),
            ("Issue Type", "One of: auto_fulfilled_by, structural_gap, credit_assumption, catalog_verify."),
            ("Group ID", "The specific requirement group within the program where the issue occurs."),
            ("Description", "Plain-language explanation of what needs to be checked or confirmed."),
            ("Resolution", "Enter your finding here once you have checked the catalog."),
            ("Resolved", "Set to Yes once addressed, No if still open, N/A if not applicable."),
        ],
        headers=["Column", "Purpose"],
        col_widths=[1.5, 4.7],
    )


def section_feedback(doc):
    heading(doc, "What Feedback Is Requested", level=1, colour=ISU_RED)
    body(doc,
        "The four issue types in the Issue Log each require a specific kind of answer. "
        "Here is what to check and what to write in the Resolution column for each."
    )

    heading(doc, "auto_fulfilled_by — Does this course satisfy the requirement?", level=2)
    body(doc,
        "Many major courses also carry gen-ed attribute codes in CourseDog, suggesting they "
        "might fulfill a general education requirement. However, there are two very different "
        "situations:"
    )
    two_col_table(doc,
        rows=[
            ("Satisfies (exempt: true)",
             "The course genuinely counts toward the gen-ed category. Example: ECO 101 is listed "
             "as fulfilling the ISU Social Sciences requirement. A student who takes ECO 101 for "
             "the major does not need to take a separate Social Sciences course.",
             "Mark Resolved: Yes. Write in Resolution: 'Satisfies — should be exempt.'"),
            ("Presupposes (auto_fulfilled_by)",
             "The course is in the same subject area as the gen-ed requirement but does not "
             "substitute for it. Example: MAT 145 covers Calculus I, which the gen-ed Mathematics "
             "requirement does not require — the major assumes the student has already met math "
             "requirements at a lower level. The course does not waive the gen-ed slot.",
             "Mark Resolved: Yes. Write in Resolution: 'Does not satisfy — keep auto_fulfilled_by.'"),
        ],
        headers=["Situation", "What it means", "What to do"],
        col_widths=[1.4, 2.8, 2.0],
    )
    callout(doc,
        "How to check:",
        "Open the program's catalog page (link in Sheet 1). Look at the program requirements "
        "and any notes about gen-ed fulfillment. If the catalog explicitly says a course "
        "satisfies a gen-ed category, it is exempt. If the course simply happens to be in a "
        "related discipline but the catalog does not say it fulfills gen-ed, it presupposes.",
        bg=LIGHT_GREY, title_colour=ISU_GREY,
    )

    heading(doc, "structural_gap — Confirm the correct program structure", level=2)
    body(doc,
        "A small number of programs have structural features that the tool could not reliably "
        "encode. These require an advisor who knows the program to describe what is correct. "
        "Examples include:"
    )
    for item in [
        "A required block that contains choice slots (e.g., pick MAT 120 or MAT 121, not both).",
        "A program divided into phases (pre-admission prerequisites, then core coursework).",
        "Applied music or ensemble requirements that allow students to repeat a course for credit over multiple semesters.",
        "Teaching specialty tracks where students choose one area of focus that determines a bundle of courses.",
    ]:
        bullet(doc, item)
    body(doc,
        "For each structural gap, the Description column explains exactly what is unknown. "
        "Write your answer in the Resolution column using plain language — a developer will "
        "translate your answer into the data file."
    )

    heading(doc, "credit_assumption — Verify course credit hours", level=2)
    body(doc,
        "In a small number of cases, the tool could not retrieve the credit hours for a course "
        "and used a standard assumption (typically 3 credit hours for a lecture course). "
        "The Description column identifies which courses were assumed. Please check the "
        "catalog page or the ISU course schedule to confirm the correct credit value and note "
        "any discrepancy in the Resolution column."
    )

    heading(doc, "catalog_verify — Confirm an interpretation", level=2)
    body(doc,
        "These items flag specific data points where two interpretations are possible and "
        "catalog language is ambiguous. Each Description explains the ambiguity. Check the "
        "catalog and — if necessary — consult the department or registrar to confirm the "
        "correct interpretation."
    )

    heading(doc, "Reporting Problems You Find During Review", level=2)
    body(doc,
        "While reviewing a worksheet, you may notice issues not pre-listed in the Issue Log. "
        "Please add a new row to Sheet 2 for any such item. You do not need to use a specific "
        "format — a clear description of the problem and where it occurs is sufficient. "
        "Mark the program's Status as 'Changes Needed' in Sheet 1."
    )

    heading(doc, "What You Do Not Need to Check", level=2)
    body(doc,
        "To keep the scope manageable, reviewers are not responsible for:"
    )
    for item in [
        "Verifying the accuracy of catalog URLs (these are auto-generated and will be tested separately).",
        "Checking gen-ed category names or descriptions — these come directly from the catalog and are not editable.",
        "Reviewing graduation requirements (120 hours, AMALI, IDEAS, etc.) — these are standardized and correct by construction.",
        "Reviewing the layout or visual design of the worksheets.",
    ]:
        bullet(doc, item)


def section_college_breakdown(doc):
    heading(doc, "Program Counts by College", level=1, colour=ISU_RED)
    body(doc,
        "The 300 program sequences are distributed across eight colleges. Review assignments "
        "should follow college lines where possible, since advisors within a college have the "
        "contextual knowledge needed to interpret catalog language for their programs."
    )
    two_col_table(doc,
        rows=[
            ("College of Arts and Sciences", "110"),
            ("Applied Science and Technology", "78"),
            ("Wonsook Kim College of Fine Arts", "50"),
            ("College of Business", "28"),
            ("College of Education", "17"),
            ("Office of the Provost", "11"),
            ("College of Engineering", "3"),
            ("Mennonite College of Nursing", "3"),
        ],
        headers=["College", "Programs"],
        col_widths=[4.5, 1.7],
    )


def section_glossary(doc):
    heading(doc, "Glossary", level=1, colour=ISU_RED)
    body(doc,
        "The Issue Log and this guide use several technical terms. Plain-language definitions follow."
    )
    two_col_table(doc,
        rows=[
            ("auto_fulfilled_by",
             "A tag on a gen-ed requirement indicating that one or more major courses might "
             "satisfy it. Needs human verification to confirm whether the course genuinely "
             "substitutes for the requirement (exempt) or merely relates to its subject area."),
            ("exempt",
             "A gen-ed slot that is automatically waived for students in this major because a "
             "major course satisfies it. The slot appears on the worksheet in a completed state."),
            ("fill type",
             "How a requirement group is satisfied. Common types: fixed (specific required "
             "courses), choose_n (pick N from a list), open (any course meeting a description), "
             "repeat (same or similar course taken multiple times for accumulating credit), "
             "escrow (credit granted upon completing specified trigger courses)."),
            ("gen-ed track",
             "The general education framework that applies to a student: ISU Gen Ed, IAI "
             "Transferable Core, or Completed Associate's Degree."),
            ("group",
             "A named block of requirements within a program, such as 'Required courses' or "
             "'Art history elective courses.' Each group has a fill type and appears as a "
             "labeled section on the worksheet."),
            ("IAI",
             "Illinois Articulation Initiative. A statewide agreement allowing students who "
             "completed the IAI General Education Core Curriculum at an Illinois community "
             "college to transfer those credits as a block satisfying university gen-ed."),
            ("JSON",
             "The data file format used to store each program's requirements. Not something "
             "reviewers need to read directly — it is the source from which worksheets are built."),
            ("phase",
             "A temporal stage within a program. For example, Nursing has a pre-nursing phase "
             "(prerequisite courses before admission) and a nursing core phase (upper-division "
             "courses after admission). Phases appear as labeled sections on the worksheet."),
            ("program sequence",
             "A named degree option within a major. For example, the Accountancy major has "
             "five sequences (Financial Accounting, Accounting Information Systems, etc.). "
             "Each sequence has its own worksheet."),
            ("slot",
             "A single row on the worksheet representing one course or one choice position "
             "within a requirement group."),
        ],
        headers=["Term", "Meaning"],
        col_widths=[1.6, 4.6],
    )


def section_contact(doc):
    heading(doc, "Questions and Submission", level=1, colour=ISU_RED)
    body(doc,
        "Please return the completed advisor-review.xlsx workbook to the Center for "
        "Integrated Professional Development (CIPD). If you have questions about the "
        "review process, the worksheets, or a specific program, contact the CIPD advising "
        "technology team."
    )
    callout(doc,
        "Important:",
        "Please do not alter the program data files or the worksheets directly. "
        "All corrections flow through the Issue Log and are applied by the development team. "
        "This ensures that both the HTML and PDF versions stay in sync.",
        bg=LIGHT_RED, title_colour=ISU_RED,
    )


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate reviewer guide Word document")
    parser.add_argument("--out", default="output/review/reviewer-guide.docx")
    args = parser.parse_args()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # Title page
    title = doc.add_heading("ISU Plans of Study Worksheets", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = ISU_RED

    sub = doc.add_paragraph("Advisor Review Guide")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = ISU_GREY

    sub2 = doc.add_paragraph("Center for Integrated Professional Development — Illinois State University")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].font.size = Pt(10)
    sub2.runs[0].font.color.rgb = ISU_GREY

    doc.add_page_break()

    # Sections
    section_introduction(doc)
    doc.add_paragraph()
    section_data_collection(doc)
    doc.add_paragraph()
    section_worksheets(doc)
    doc.add_paragraph()
    section_workbook(doc)
    doc.add_paragraph()
    section_feedback(doc)
    doc.add_paragraph()
    section_college_breakdown(doc)
    doc.add_paragraph()
    section_glossary(doc)
    doc.add_paragraph()
    section_contact(doc)

    os.makedirs(os.path.dirname(args.out), exist_ok=True)
    doc.save(args.out)
    print(f"Wrote {args.out}")


if __name__ == "__main__":
    main()
